
import sys

# Improting Image class from PIL module
from docx import Document
from docx.shared import Inches
from PIL import Image
import fitz

pdffile = "infile.pdf"
doc = fitz.open(pdffile)
page = doc.loadPage(0)  # number of page
pix = page.getPixmap()
output = "outfile.png"
pix.writePNG(output)


# Opens a image in RGB mode
im = Image.open(r"C:\Users\Admin\Pictures\geeks.png")

# Size of the image in pixels (size of orginal image)
# (This is not mandatory)
width, height = im.size

# Setting the points for cropped image
left = 5
top = height / 4
right = 164
bottom = 3 * height / 4

# Cropped image of above dimension
# (It will not change orginal image)
im1 = im.crop((left, top, right, bottom))

# Shows the image in image viewer
im1.show()


#### pictures to word file


document = Document()

document.add_heading('Document Title', 0)

document.add_picture(image_filepath, width=Inches(1.25))
document.add_picture(image_filepath, width=Inches(1.25))

document.save(doc_filepath)