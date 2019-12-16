
import sys
# Improting Image class from PIL module
from docx import Document
from docx.shared import Inches
from PIL import Image
import fitz


print("enter file path")
pdffile=input()
doc = fitz.open(pdffile)
print("enter pages num in file")
pages=input()
page = doc.loadPage(int(pages))  # number of page
pix = page.getPixmap()
output = "outfile.png"
pix.writePNG(output)

# Opens a image in RGB mode
print("enter path for saved photo")
path=input()
im = Image.open("r"+""+path)

# Size of the image in pixels (size of orginal image)
# (This is not mandatory)
width, height = im.size

# Setting the points for cropped image
print("enter left size")
leftsize=input()
left = int(leftsize)
print("enter top size")
topsize=input()
top = int(topsize) / 4
print("enter right size")
rightsize=input()
right = int(rightsize)
print("enter bottom size")
bottomsize=input()
bottom = 3 * int(bottomsize) / 4

# Cropped image of above dimension
# (It will not change orginal image)
im1 = im.crop((left, top, right, bottom))

# Shows the image in image viewer
im1.show()


#### pictures to word file


document = Document()

document.add_heading('Document Title', 0)
print("enter cropped image file path")
image_filepath= input()
document.add_picture(image_filepath, width=Inches(1.25))
document.add_picture(image_filepath, width=Inches(1.25))

print("enter doc file path")
doc_filepath= input()
document.save(doc_filepath)